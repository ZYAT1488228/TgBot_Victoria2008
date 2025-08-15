from docx import Document
from datetime import date
from docx.shared import Pt
from io import BytesIO
import os

def make_pass(
    names: list[str],
    birth_dates: list[str],
    passport_numbers: list[str],
    start_date: date,
    end_date: date,
    created_at: date,
    auto_model: str,
    auto_plates: str,
    num: str,
    passport_issued_by: str = '',
    passport_id_code: str = ''
):
    # Абсолютный путь к шаблону, чтобы не зависеть от рабочей директории (systemd и т.п.)
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(BASE_DIR, '..', 'templates', 'template.docx')

    doc = Document(template_path)

    start_date = start_date.strftime("%d.%m.%Y")
    end_date = end_date.strftime("%d.%m.%Y")
    created_at = created_at.strftime("%d.%m.%Y")

    person_data = ""
    for i in range(len(names)):
        person_data += f"ПІБ: {names[i]}\n"
        person_data += f"Дата народження: {birth_dates[i]}\n"
        person_data += f"№ Паспорта: {passport_numbers[i]}\n\n"

    # Заменяем плейсхолдер <person_data>
    for paragraph in doc.paragraphs:
        if '<person_data>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<person_data>', person_data.strip())

    # Заменяем остальные плейсхолдеры
    for paragraph in doc.paragraphs:
        if '<date_start>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<date_start>', start_date)
        if '<date_end>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<date_end>', end_date)
        if '<todaysdate>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<todaysdate>', created_at)
        if '<automodel>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<automodel>', f'Авто: {auto_model}' if auto_model else '')
        if '<autoplates>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<autoplates>', f'Гос.Номер: {auto_plates}' if auto_plates else '')
        if '<number>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<number>', num)
        if '<passport_issued_by>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<passport_issued_by>', passport_issued_by)
        if '<passport_id_code>' in paragraph.text:
            paragraph.text = paragraph.text.replace('<passport_id_code>', passport_id_code)

        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'

    # Безопасное имя файла (убираем пробелы/кириллицу/спецсимволы, оставляем A-Z a-z 0-9 _ -)
    base_name = "".join(c for c in names[0] if c.isalnum() or c in ('_', '-'))
    new_word_file = f'{num}_{base_name}_{start_date}_{end_date}.docx'

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # критично: перемотать поток перед отправкой в Telegram

    return file_stream, new_word_file
